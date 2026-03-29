"""
Генератор DOCX-руководства ПРИЗМА.
Вместо текстовых таблиц — скриншоты с аннотациями и выносные фрагменты.
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOCS_DIR = Path(__file__).parent
SCREENSHOTS = DOCS_DIR / "screenshots"
CROPS = SCREENSHOTS / "crops"
OUT_FILE = DOCS_DIR / "USER_GUIDE.docx"

# Colors
C_PRIMARY = RGBColor(0x1A, 0x73, 0xE8)
C_DARK = RGBColor(0x1E, 0x29, 0x3B)
C_GRAY = RGBColor(0x5F, 0x63, 0x68)
C_ACCENT = RGBColor(0x34, 0xA8, 0x53)
TBL_HDR = "1A73E8"
TBL_ALT = "F0F4F9"


def shading(el, color):
    el.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))


def img(doc, path, caption=None, width=Inches(6.2)):
    """Insert image centered with optional caption."""
    if not path.exists():
        p = doc.add_paragraph()
        r = p.add_run(f"[{path.name} — файл не найден]")
        r.font.italic = True
        r.font.color.rgb = RGBColor(0xFF, 0, 0)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=width)
    if caption:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_after = Pt(8)
        r = pc.add_run(caption)
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = C_GRAY


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text, bold=False, italic=False, size=Pt(10.5), color=None, align=None,
         space_before=None, space_after=None, indent=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_before:
        p.paragraph_format.space_before = space_before
    if space_after:
        p.paragraph_format.space_after = space_after
    if indent:
        p.paragraph_format.left_indent = indent
    r = p.add_run(text)
    r.font.size = size
    r.font.name = 'Calibri'
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def rich_para(doc, parts, align=None, space_before=None, space_after=None, indent=None):
    """parts = list of (text, {bold, italic, color, size, font})"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_before:
        p.paragraph_format.space_before = space_before
    if space_after:
        p.paragraph_format.space_after = space_after
    if indent:
        p.paragraph_format.left_indent = indent
    for text, fmt in parts:
        r = p.add_run(text)
        r.font.name = fmt.get('font', 'Calibri')
        r.font.size = fmt.get('size', Pt(10.5))
        r.font.bold = fmt.get('bold', False)
        r.font.italic = fmt.get('italic', False)
        if 'color' in fmt:
            r.font.color.rgb = fmt['color']
    return p


def callout(doc, text, kind="tip"):
    """Colored callout box."""
    colors = {"tip": "E8F5E9", "warning": "FFF3E0", "important": "FCE4EC"}
    icons = {"tip": "💡 Совет.", "warning": "⚠ Внимание.", "important": "❗ Важно."}
    p = doc.add_paragraph()
    shading(p._p.get_or_add_pPr(), colors.get(kind, "E8F5E9"))
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    ri = p.add_run(icons.get(kind, "💡") + " ")
    ri.font.bold = True
    ri.font.size = Pt(10)
    ri.font.name = 'Calibri'
    rt = p.add_run(text)
    rt.font.size = Pt(10)
    rt.font.name = 'Calibri'


def bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if bold_prefix:
        r = p.add_run(bold_prefix + " — ")
        r.font.bold = True
        r.font.size = Pt(10.5)
        r.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = 'Calibri'


def numbered(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    rn = p.add_run(f"{num}. ")
    rn.font.bold = True
    rn.font.color.rgb = C_PRIMARY
    rn.font.size = Pt(10.5)
    rn.font.name = 'Calibri'
    rt = p.add_run(text)
    rt.font.size = Pt(10.5)
    rt.font.name = 'Calibri'


def small_table(doc, headers, rows):
    """Small reference table (for glossary, hotkeys, etc.)."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shading(c._tc.get_or_add_tcPr(), TBL_HDR)
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9.5)
                r.font.name = 'Calibri'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = val
            if ri % 2 == 1:
                shading(c._tc.get_or_add_tcPr(), TBL_ALT)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    r.font.name = 'Calibri'
    doc.add_paragraph()


def separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("━" * 50)
    r.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    r.font.size = Pt(8)


def configure_styles(doc):
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(10.5)
    s.font.color.rgb = C_DARK
    s.paragraph_format.space_before = Pt(2)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.15

    for lv, (sz, col) in enumerate([
        (Pt(22), C_PRIMARY), (Pt(16), C_DARK), (Pt(13), C_DARK)
    ], 1):
        h = doc.styles[f'Heading {lv}']
        h.font.name = 'Calibri'
        h.font.size = sz
        h.font.bold = True
        h.font.color.rgb = col
        h.paragraph_format.space_before = Pt(18 if lv == 1 else 12)
        h.paragraph_format.space_after = Pt(6)
        if lv == 1:
            h.paragraph_format.page_break_before = True


def cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    para(doc, "ПРИЗМА", bold=True, size=Pt(48), color=C_PRIMARY,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Планирование Ресурсов И Задач, Мониторинг и Анализ",
         italic=True, size=Pt(14), color=C_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    para(doc, "━" * 40, size=Pt(14), color=C_PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    para(doc, "Руководство пользователя", bold=True, size=Pt(24), color=C_DARK,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    para(doc, "Версия 1.0  |  29.03.2026", size=Pt(12), color=C_GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def toc(doc):
    # No page break for TOC heading
    h = doc.add_heading("Содержание", level=1)
    h.paragraph_format.page_break_before = False

    sections = [
        ("1", "Начало работы"),
        ("2", "Навигация по системе"),
        ("3", "Стартовая страница (Дашборд)"),
        ("4", "Управление проектами (УП)"),
        ("5", "Производственный план (ПП)"),
        ("6", "Сводное планирование (СП)"),
        ("7", "Журнал извещений (ЖИ)"),
        ("8", "Аналитика"),
        ("9", "Производственный календарь"),
        ("10", "Сотрудники"),
        ("11", "План отпусков и командировок"),
        ("12", "Профиль и настройки"),
        ("13", "Горячие клавиши и быстрые действия"),
        ("14", "Роли и права доступа"),
        ("15", "Типовые сценарии работы"),
        ("16", "Устранение неполадок и FAQ"),
        ("А", "Глоссарий"),
    ]
    for num, title in sections:
        rich_para(doc, [
            (f"  {num}  ", {"bold": True, "color": C_PRIMARY, "size": Pt(11)}),
            (title, {"size": Pt(11), "color": C_DARK}),
        ], space_before=Pt(2), space_after=Pt(2))
    doc.add_page_break()


def add_footer(doc):
    sect = doc.sections[0]
    ft = sect.footer
    ft.is_linked_to_previous = False
    fp = ft.paragraphs[0] if ft.paragraphs else ft.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = fp.add_run("━" * 30 + "\n")
    rl.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    rl.font.size = Pt(7)
    rf = fp.add_run("ПРИЗМА — Руководство пользователя  |  v1.0  |  ")
    rf.font.size = Pt(8)
    rf.font.color.rgb = C_GRAY
    rf.font.name = 'Calibri'
    # Page number
    for xml_str in [
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>',
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>',
        f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>',
    ]:
        r = fp.add_run()
        r._r.append(parse_xml(xml_str))
    r2 = fp.add_run("1")
    r2.font.size = Pt(8)
    r2.font.color.rgb = C_GRAY
    r3 = fp.add_run()
    r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


def build():
    doc = Document()
    sect = doc.sections[0]
    sect.page_width = Cm(21)
    sect.page_height = Cm(29.7)
    sect.top_margin = Cm(2)
    sect.bottom_margin = Cm(2)
    sect.left_margin = Cm(2.5)
    sect.right_margin = Cm(2)

    configure_styles(doc)
    cover(doc)
    toc(doc)

    # ========== 1. НАЧАЛО РАБОТЫ ==========
    heading(doc, "1. Начало работы")

    heading(doc, "1.1 Вход в систему", 2)
    numbered(doc, 1, "Откройте браузер и перейдите по адресу системы.")
    numbered(doc, 2, "На экране авторизации введите логин и пароль.")
    numbered(doc, 3, 'Нажмите кнопку «Войти».')

    img(doc, SCREENSHOTS / "login.png", "Рис. 1.1. Экран авторизации ПРИЗМА")

    para(doc, "Выносной элемент — форма входа крупным планом:", bold=True,
         space_before=Pt(6))
    img(doc, CROPS / "login_form.png", "Рис. 1.2. Форма входа (увеличено)", width=Inches(4))

    callout(doc, 'Если вы забыли пароль, нажмите ссылку «Забыли пароль?» под формой входа.', "tip")

    heading(doc, "1.2 Первый вход и обучение", 2)
    para(doc, "При первом входе запускается интерактивный тур — 17 шагов по всем модулям системы.")

    img(doc, SCREENSHOTS / "tour_step.png", "Рис. 1.3. Обучающий тур — подсветка элемента")

    para(doc, "Выносной элемент — карточка тура:", bold=True, space_before=Pt(6))
    img(doc, CROPS / "tour_card.png", "Рис. 1.4. Информационная карточка тура (увеличено)", width=Inches(4.5))

    para(doc, "Управление туром:")
    bullet(doc, "кнопка «Далее» или клавиша Enter", bold_prefix="Следующий шаг")
    bullet(doc, "клавиша Esc", bold_prefix="Пропустить тур")
    bullet(doc, "Профиль → «Пройти обучение»", bold_prefix="Перезапустить тур")

    separator(doc)

    # ========== 2. НАВИГАЦИЯ ==========
    heading(doc, "2. Навигация по системе")

    heading(doc, "2.1 Боковая панель (сайдбар)", 2)
    para(doc, "Левая боковая панель — основной способ навигации. Всегда видна на экране.")

    img(doc, SCREENSHOTS / "sidebar.png", "Рис. 2.1. Боковая панель навигации", width=Inches(1.5))

    para(doc, "Основные разделы:", bold=True)
    bullet(doc, "Дашборд — персональная сводка, метрики", bold_prefix="🏠 Стартовая страница")
    bullet(doc, "Реестр проектов и изделий", bold_prefix="📁 Управление проектами")
    bullet(doc, "Формирование и ведение ПП подразделения", bold_prefix="🏭 Производственный план")
    bullet(doc, "Единый план-отчёт с помесячными часами", bold_prefix="📋 Сводное планирование")
    bullet(doc, "Учёт ИИ/ПИ с контролем сроков погашения", bold_prefix="📰 Журнал извещений")
    bullet(doc, "Графики загрузки, исполнения, метрики", bold_prefix="📊 Аналитика")

    callout(doc, 'Нажмите кнопку ☰ в шапке, чтобы свернуть сайдбар.', "tip")

    heading(doc, "2.2 Командная палитра (Ctrl+K)", 2)
    para(doc, "Быстрый поиск по разделам и действиям — нажмите Ctrl+K на любой странице.")

    img(doc, SCREENSHOTS / "command_palette.png", "Рис. 2.2. Командная палитра")

    para(doc, "Выносной элемент — окно палитры крупным планом:", bold=True, space_before=Pt(6))
    img(doc, CROPS / "cmd_zoom.png", "Рис. 2.3. Командная палитра (увеличено)", width=Inches(4.5))

    separator(doc)

    # ========== 3. ДАШБОРД ==========
    heading(doc, "3. Стартовая страница (Дашборд)")

    heading(doc, "3.1 Общий вид", 2)
    para(doc, "Стартовая страница — персональная сводка с метриками загрузки.")

    img(doc, SCREENSHOTS / "dashboard.png", "Рис. 3.1. Стартовая страница (полный вид)")

    para(doc, "Аннотированный скриншот — основные области:", bold=True, space_before=Pt(6))
    img(doc, CROPS / "dash_annotated.png", "Рис. 3.2. Дашборд — области экрана с аннотациями")

    heading(doc, "3.2 Карточки показателей (увеличено)", 2)
    img(doc, CROPS / "dash_metrics.png",
        "Рис. 3.3. Карточки: Загрузка %, План/Норма, Выполнено, В работе, Долги, Ср. просрочка")

    heading(doc, "3.3 Блок «Команда» (увеличено)", 2)
    img(doc, CROPS / "dash_team.png", "Рис. 3.4. Список отделов с показателями загрузки")

    para(doc, "Цветовое кодирование загрузки:", bold=True)
    bullet(doc, "зелёный (недозагрузка)", bold_prefix="< 70%")
    bullet(doc, "обычный (нормальная загрузка)", bold_prefix="70–100%")
    bullet(doc, "красный (перегрузка)", bold_prefix="> 100%")

    separator(doc)

    # ========== 4. УПРАВЛЕНИЕ ПРОЕКТАМИ ==========
    heading(doc, "4. Управление проектами (УП)")

    callout(doc, "Доступ: только администраторы.", "warning")

    heading(doc, "4.1 Обзор страницы", 2)
    img(doc, SCREENSHOTS / "projects.png", "Рис. 4.1. Карточная сетка проектов")

    para(doc, "Выносной элемент — карточка проекта:", bold=True, space_before=Pt(6))
    img(doc, CROPS / "proj_card.png", "Рис. 4.2. Карточка проекта (увеличено)", width=Inches(3))

    heading(doc, "4.2 Создание проекта", 2)
    numbered(doc, 1, 'Нажмите «Добавить проект».')
    numbered(doc, 2, "Заполните: полное наименование (обязательно), краткое наименование, шифр.")
    numbered(doc, 3, 'Нажмите «Сохранить».')

    callout(doc, "Краткое наименование используется для автогенерации кодов строк ПП (формат: Орбита-НП.1, Орбита-НП.2, ...).", "tip")

    heading(doc, "4.3 Управление изделиями", 2)
    para(doc, "Каждый проект содержит список изделий. Они отображаются как цветные бейджи на карточке проекта.")

    separator(doc)

    # ========== 5. ПРОИЗВОДСТВЕННЫЙ ПЛАН ==========
    heading(doc, "5. Производственный план (ПП)")

    heading(doc, "5.1 Выбор проекта ПП", 2)
    img(doc, SCREENSHOTS / "pp_projects.png", "Рис. 5.1. Список проектов производственного плана")

    heading(doc, "5.2 Таблица ПП — полный вид", 2)
    img(doc, SCREENSHOTS / "pp_table.png", "Рис. 5.2. Таблица ПП (полный экран)")

    heading(doc, "5.3 Аннотированная структура ПП", 2)
    para(doc, "На скриншоте отмечены три ключевые области таблицы ПП:")
    img(doc, CROPS / "pp_table_annotated.png",
        "Рис. 5.3. Структура ПП: ① Статусная панель, ② Тулбар, ③ Заголовки + фильтры")

    heading(doc, "5.4 Статусная панель (увеличено)", 2)
    img(doc, CROPS / "pp_status_panel.png",
        "Рис. 5.4. Статусная панель — Все / Выполнено / Просрочено / В работе")
    para(doc, "Нажмите на любой счётчик для фильтрации таблицы по этому статусу.")

    heading(doc, "5.5 Панель инструментов (увеличено)", 2)
    img(doc, CROPS / "pp_toolbar.png",
        "Рис. 5.5. Панель инструментов: поиск, добавление строки, синхронизация")

    heading(doc, "5.6 Заголовки и фильтры (увеличено)", 2)
    img(doc, CROPS / "pp_headers.png",
        "Рис. 5.6. Заголовки столбцов с кнопками мульти-фильтров (▼)")

    heading(doc, "5.7 Строки данных (увеличено)", 2)
    img(doc, CROPS / "pp_rows_zoom.png",
        "Рис. 5.7. Строки ПП — двойной клик по ячейке для inline-редактирования")

    heading(doc, "5.8 Inline-редактирование", 2)
    numbered(doc, 1, "Дважды кликните по ячейке.")
    numbered(doc, 2, "Ячейка превращается в поле ввода.")
    numbered(doc, 3, "Enter — сохранить, Esc — отмена.")

    callout(doc, "Вы можете редактировать только строки своего подразделения.", "warning")

    heading(doc, "5.9 Зависимости задач", 2)
    small_table(doc,
        ["Тип", "Название", "Описание"],
        [
            ["FS", "Finish-to-Start", "B начинается после завершения A"],
            ["SS", "Start-to-Start", "B начинается вместе с A"],
            ["FF", "Finish-to-Finish", "B завершается вместе с A"],
            ["SF", "Start-to-Finish", "B завершается после начала A"],
        ])

    callout(doc, "Система проверяет циклы (A→B→C→A). Зависимость с циклом не сохранится.", "important")

    heading(doc, "5.10 Синхронизация ПП → СП", 2)
    para(doc, 'Кнопка «⇄ Синхронизировать с СП» переносит строки ПП в сводное планирование.')
    para(doc, "После синхронизации в задачах СП блокируются поля: наименование, номер, обозначение, этап, обоснование (🔒).")

    separator(doc)

    # ========== 6. СВОДНОЕ ПЛАНИРОВАНИЕ ==========
    heading(doc, "6. Сводное планирование (СП)")

    heading(doc, "6.1 Общий вид", 2)
    img(doc, SCREENSHOTS / "plan_table.png", "Рис. 6.1. Таблица сводного планирования (полный вид)")

    heading(doc, "6.2 Аннотированная структура СП", 2)
    img(doc, CROPS / "sp_table_annotated.png",
        "Рис. 6.2. Структура СП: ① Статусная панель, ② Тулбар, ③ Заголовки")

    heading(doc, "6.3 Строки СП с помесячными часами (увеличено)", 2)
    img(doc, CROPS / "sp_rows_zoom.png",
        "Рис. 6.3. Строки СП — столбцы Янв–Дек с плановыми часами")

    heading(doc, "6.4 Создание задачи", 2)
    numbered(doc, 1, 'Нажмите «+ Задача» или Ctrl+K → «Новая задача».')
    numbered(doc, 2, "В модальном окне заполните: наименование, проект, исполнитель, даты, помесячные часы.")
    numbered(doc, 3, 'Нажмите «Сохранить».')

    heading(doc, "6.5 Ошибки планирования", 2)
    para(doc, 'Кнопка «⚠ Ошибки» проверяет:')
    bullet(doc, "задачи без исполнителя")
    bullet(doc, "задачи без часов")
    bullet(doc, "пересечения с отпуском исполнителя")
    bullet(doc, "задачи без дат")

    heading(doc, "6.6 Управление столбцами и плотность", 2)
    para(doc, 'Кнопка «⚙ Столбцы» — показать/скрыть, менять порядок. Ширины сохраняются на сервере.')
    para(doc, "Три режима плотности строк: компактный, средний, просторный.")

    separator(doc)

    # ========== 7. ЖУРНАЛ ИЗВЕЩЕНИЙ ==========
    heading(doc, "7. Журнал извещений (ЖИ)")

    heading(doc, "7.1 Общий вид", 2)
    img(doc, SCREENSHOTS / "notices.png", "Рис. 7.1. Журнал извещений (полный вид)")

    heading(doc, "7.2 Аннотированная структура", 2)
    img(doc, CROPS / "notices_annotated.png",
        "Рис. 7.2. Структура ЖИ: заголовки столбцов и цветовые статусы")

    heading(doc, "7.3 Строки журнала (увеличено)", 2)
    img(doc, CROPS / "notices_rows.png", "Рис. 7.3. Строки журнала извещений (увеличено)")

    heading(doc, "7.4 Типы и статусы", 2)
    small_table(doc, ["Тип", "Описание"], [
        ["ИИ", "Извещение об изменении — постоянное"],
        ["ПИ", "Предварительное извещение — временное"],
    ])
    small_table(doc, ["Статус", "Цвет", "Описание"], [
        ["Активное", "🟢", "ИИ выпущено, ожидает погашения"],
        ["Истекшее", "🟡", "ПИ с прошедшей датой действия"],
        ["Погашено (да)", "🔵", "Выпущено погашающее извещение"],
        ["Погашено (нет)", "⚪", "Закрыто без погашающего извещения"],
    ])

    callout(doc, "При типе работы «Корректировка документа» запись в ЖИ создаётся автоматически.", "tip")

    separator(doc)

    # ========== 8. АНАЛИТИКА ==========
    heading(doc, "8. Аналитика")

    heading(doc, "8.1 Общий вид", 2)
    img(doc, SCREENSHOTS / "analytics.png", "Рис. 8.1. Модуль аналитики (полный вид)")

    heading(doc, "8.2 Аннотированная структура", 2)
    img(doc, CROPS / "analytics_annotated.png",
        "Рис. 8.2. Аналитика: ① Карточки-итоги, ② Графики, ③ Таблица по отделам")

    heading(doc, "8.3 Графики загрузки (увеличено)", 2)
    img(doc, CROPS / "analytics_chart.png",
        "Рис. 8.3. Столбчатая диаграмма загрузки по отделам и месяцам")

    para(doc, "Три аналитических доски:", bold=True)
    bullet(doc, "загрузка по НТЦ и отделам, дедлайны, drilldown по сотрудникам",
           bold_prefix="Доска руководителя")
    bullet(doc, "персональная загрузка, список задач",
           bold_prefix="Доска сотрудника")
    bullet(doc, "% выполнения, трудоёмкость по этапам",
           bold_prefix="Доска ПП")

    separator(doc)

    # ========== 9. ПРОИЗВОДСТВЕННЫЙ КАЛЕНДАРЬ ==========
    heading(doc, "9. Производственный календарь")

    heading(doc, "9.1 Общий вид", 2)
    img(doc, SCREENSHOTS / "work_calendar.png", "Рис. 9.1. Производственный календарь")

    heading(doc, "9.2 Аннотированная структура", 2)
    img(doc, CROPS / "cal_annotated.png",
        "Рис. 9.2. Календарь: ① Сводная таблица часов, ② Мини-календари")

    heading(doc, "9.3 Мини-календари (увеличено)", 2)
    img(doc, CROPS / "cal_mini.png",
        "Рис. 9.3. Мини-календари: выходные (красный) и праздники")

    callout(doc, "Праздничные дни автоматически учитываются при выравнивании дат зависимостей.", "tip")

    separator(doc)

    # ========== 10. СОТРУДНИКИ ==========
    heading(doc, "10. Сотрудники")

    heading(doc, "10.1 Общий вид", 2)
    img(doc, SCREENSHOTS / "employees.png", "Рис. 10.1. Список сотрудников")

    heading(doc, "10.2 Аннотированный экран", 2)
    img(doc, CROPS / "emp_annotated.png",
        "Рис. 10.2. Поиск и фильтрация сотрудников (с аннотацией)")

    heading(doc, "10.3 Таблица с ролями (увеличено)", 2)
    img(doc, CROPS / "emp_roles.png",
        "Рис. 10.3. Таблица: ФИО, роль (цветной бейдж), должность, отдел, статус")

    small_table(doc, ["Цвет", "Роль"], [
        ["🔴", "Администратор"],
        ["🔵", "Руководитель НТЦ / Зам. руководителя"],
        ["🟡", "Начальник отдела / Зам. начальника"],
        ["🟢", "Начальник сектора"],
        ["⚪", "Исполнитель"],
    ])

    separator(doc)

    # ========== 11. ПЛАН ОТПУСКОВ ==========
    heading(doc, "11. План отпусков и командировок")
    para(doc, "Таблицы отпусков и командировок с фильтрацией по подразделению.")
    callout(doc, "Отпуска учитываются при проверке ошибок планирования в СП.", "tip")

    separator(doc)

    # ========== 12. ПРОФИЛЬ ==========
    heading(doc, "12. Профиль и настройки")

    heading(doc, "12.1 Общий вид", 2)
    img(doc, SCREENSHOTS / "profile.png", "Рис. 12.1. Страница профиля")

    heading(doc, "12.2 Аннотированная структура", 2)
    img(doc, CROPS / "profile_annotated.png",
        "Рис. 12.2. Профиль: ① Карточка пользователя, ② Моё меню")

    heading(doc, "12.3 Моё меню (увеличено)", 2)
    img(doc, CROPS / "profile_menu.png",
        "Рис. 12.3. Настройки: сброс ширин, видимость отделов, тема, обучение", width=Inches(4))

    para(doc, "Темы оформления:", bold=True)
    bullet(doc, "белый фон — для светлых помещений", bold_prefix="☀ Светлая")
    bullet(doc, "приглушённые тона", bold_prefix="🌤 Сумеречная")
    bullet(doc, "тёмный фон — снижает нагрузку на глаза", bold_prefix="🌙 Тёмная")
    bullet(doc, "автоматически следует настройке ОС", bold_prefix="🖥 Системная")

    separator(doc)

    # ========== 13. ГОРЯЧИЕ КЛАВИШИ ==========
    heading(doc, "13. Горячие клавиши и быстрые действия")

    small_table(doc, ["Клавиша", "Действие", "Где работает"], [
        ["Ctrl+K / Cmd+K", "Командная палитра", "Везде"],
        ["Esc", "Закрыть модальное окно", "Везде"],
        ["Esc", "Пропустить шаг обучения", "Тур"],
        ["Enter", "Следующий шаг обучения", "Тур"],
        ["Двойной клик", "Inline-редактирование ячейки", "ПП"],
        ["Enter (в ячейке)", "Сохранить изменения", "ПП"],
        ["Esc (в ячейке)", "Отменить изменения", "ПП"],
    ])

    separator(doc)

    # ========== 14. РОЛИ ==========
    heading(doc, "14. Роли и права доступа")

    small_table(doc, ["Роль", "Создание", "Редактирование", "Удаление", "Админ"], [
        ["Администратор", "Все данные", "Все данные", "Все данные", "Полный"],
        ["Руководитель НТЦ", "Свой НТЦ", "Свой НТЦ", "Свой НТЦ", "—"],
        ["Зам. рук. НТЦ", "Свой НТЦ", "Свой НТЦ", "Свой НТЦ", "—"],
        ["Начальник отдела", "Свой отдел", "Свой отдел", "Свой отдел", "—"],
        ["Зам. нач. отдела", "Свой отдел", "Свой отдел", "Свой отдел", "—"],
        ["Начальник сектора", "Свой сектор", "Свой сектор", "Свой сектор", "—"],
        ["Исполнитель", "—", "—", "—", "—"],
    ])

    heading(doc, "14.2 Область видимости", 2)
    para(doc, "Иерархия: Администратор → Руководитель НТЦ → Нач. отдела → Нач. сектора → Исполнитель.")
    para(doc, "Каждый уровень видит данные только своего подразделения и ниже.")

    callout(doc, "ПП — общий документ, видимый всем (без фильтрации по ролям).", "tip")

    separator(doc)

    # ========== 15. ТИПОВЫЕ СЦЕНАРИИ ==========
    heading(doc, "15. Типовые сценарии работы")

    heading(doc, "15.1 Начальник отдела формирует план", 2)
    numbered(doc, 1, "Перейти в «Производственный план» → выбрать проект")
    numbered(doc, 2, "Добавить строки работ (кнопка «+»)")
    numbered(doc, 3, "Заполнить: наименование, исполнитель, дата, норма, коэффициент")
    numbered(doc, 4, "Создать зависимости, выровнять даты")
    numbered(doc, 5, "Синхронизировать с СП (кнопка «⇄»)")
    numbered(doc, 6, "В СП распределить часы по месяцам")
    numbered(doc, 7, "Проверить ошибки (кнопка «⚠»)")

    heading(doc, "15.2 Исполнитель просматривает задачи", 2)
    numbered(doc, 1, "Войти → Дашборд (текущие задачи, уведомления)")
    numbered(doc, 2, "Перейти в «Сводное планирование»")
    numbered(doc, 3, "Открыть задачу (✏) → посмотреть детали")
    numbered(doc, 4, "Создать отчёт (📊) после выполнения")

    heading(doc, "15.3 Руководитель НТЦ контролирует загрузку", 2)
    numbered(doc, 1, "Дашборд → отделы с загрузкой > 100% (красный)")
    numbered(doc, 2, "Аналитика → «Доска руководителя» → drilldown по отделу")
    numbered(doc, 3, "СП → фильтр по отделу и статусу «Просрочено»")
    numbered(doc, 4, "Перераспределить задачи / скорректировать сроки")

    separator(doc)

    # ========== 16. FAQ ==========
    heading(doc, "16. Устранение неполадок и FAQ")

    for q, a in [
        ("Как быстро перейти на нужную страницу?",
         "Ctrl+K → начните вводить название."),
        ("Как сменить тему?",
         "Профиль → «Моё меню» → выбор темы. Или Ctrl+K → «Сменить тему»."),
        ("Почему не могу редактировать строку ПП?",
         "1) Строка другого подразделения; 2) Роль «Исполнитель»; 3) Делегирование истекло."),
        ("Как перенести данные из ПП в СП?",
         "Кнопка «⇄ Синхронизировать с СП» на странице ПП."),
        ("Почему поля задачи заблокированы (🔒)?",
         "Задача из ПП. Редактируйте в ПП и повторите синхронизацию."),
        ("Номер извещения уже существует?",
         "Система проверяет уникальность. Используйте другой номер."),
        ("Как сбросить ширины столбцов?",
         "Профиль → «Моё меню» → «Сбросить ширины столбцов»."),
    ]:
        rich_para(doc, [
            ("В: ", {"bold": True, "color": C_PRIMARY}),
            (q, {"bold": True, "color": C_PRIMARY}),
        ], space_before=Pt(8))
        rich_para(doc, [
            ("О: ", {"bold": True}),
            (a, {}),
        ], indent=Cm(0.5))

    separator(doc)

    # ========== ПРИЛОЖЕНИЕ А ==========
    heading(doc, "Приложение А. Глоссарий")

    small_table(doc, ["Термин", "Описание"], [
        ["ПП", "Производственный план — план работ по конкретному проекту"],
        ["СП", "Сводное планирование — единый план-отчёт подразделения"],
        ["ЖИ", "Журнал извещений — учёт ИИ и ПИ"],
        ["ИИ", "Извещение об изменении"],
        ["ПИ", "Предварительное извещение"],
        ["УП", "Управление проектами — реестр проектов и изделий"],
        ["НТЦ", "Научно-технический центр"],
        ["FS", "Finish-to-Start — тип зависимости задач"],
        ["Inline-ред.", "Редактирование ячейки «на месте» двойным кликом"],
        ["Дашборд", "Стартовая страница с персональной сводкой"],
    ])

    add_footer(doc)
    doc.save(str(OUT_FILE))
    print(f"OK: {OUT_FILE} ({OUT_FILE.stat().st_size // 1024} KB)")


if __name__ == '__main__':
    build()
